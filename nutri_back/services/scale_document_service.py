from datetime import datetime
from pathlib import Path
import re
from typing import Any, Dict

from docx import Document
from docx.oxml.ns import qn

from services import sga_service


SCALES_DIR = Path(__file__).resolve().parents[1] / "scales"
OUTPUT_DIR = SCALES_DIR / "outputs"
REPORT_FILE_LABELS = {
    "NRS-2002": "NRS-2002",
    "MNA-SF": "MNA-SF",
    "GLIM": "GLIM",
    "SGA": "面部图像筛查(SGA)",
}

TEMPLATE_FILES = {
    "nrs": SCALES_DIR / "NRS-2002 评估表.docx",
    "mna": SCALES_DIR / "MNA-SF简表.docx",
    "glim": SCALES_DIR / "GLIM评估表.docx",
    "sga": SCALES_DIR / "SGA评估表.docx",
}

GI_LABELS = {
    "dysphagia": "吞咽困难", "nausea_vomiting": "恶心，呕吐",
    "diarrhea": "腹泻", "constipation": "便秘", "abdominal_pain": "腹痛", "other": "其他",
}
NUTRITION_IMPACT_LABELS = {
    "short_bowel": "短肠综合征", "pancreatic_insufficiency": "胰腺功能不全",
    "post_bariatric": "减肥手术后", "esophageal_stricture": "食管狭窄",
    "gastroparesis": "胃轻瘫", "intestinal_obstruction": "肠梗阻",
    "diarrhea_or_steatorrhea": "腹泻或脂肪痢", "high_output_stoma": "排出量较大的肠道造口患者", "other": "其他",
}
ACUTE_DISEASE_LABELS = {
    "severe_infection": "严重感染", "burn": "烧伤", "trauma": "创伤", "brain_injury": "闭合性脑损伤",
}
CHRONIC_DISEASE_LABELS = {
    "malignant_tumor": "恶性肿瘤（癌症）", "copd": "慢性阻塞性肺疾病", "heart_failure": "充血性心衰",
    "ckd": "慢性肾脏疾病", "chronic_liver": "慢性肝病", "liver_cirrhosis": "慢性肝病",
    "rheumatoid_arthritis": "类风湿性关节炎",
}


def _value(payload: Any, name: str, default: Any = "") -> Any:
    if isinstance(payload, dict):
        return payload.get(name, default)
    return getattr(payload, name, default)


def _safe_filename_part(value: str) -> str:
    cleaned = re.sub(r'[\/:*?"<>|\s]+', '_', str(value or '未命名患者')).strip('._')
    return cleaned[:40] or "未命名患者"


def _new_output_path(scale: str, payload: Any) -> Path:
    """按“日期_时间_患者_量表评估表.docx”生成可归档的报告名称。"""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    # 同一秒重复生成时自动追加序号，避免同名覆盖。
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    patient_name = _safe_filename_part(_value(payload, "patient_name"))
    scale_label = REPORT_FILE_LABELS.get(scale, scale)
    base_name = f"{timestamp}_{patient_name}_{scale_label}评估表"
    output_path = OUTPUT_DIR / f"{base_name}.docx"
    sequence = 2
    while output_path.exists():
        output_path = OUTPUT_DIR / f"{base_name}_{sequence}.docx"
        sequence += 1
    return output_path


def _check_word_checkbox(cell: Any) -> bool:
    for control in cell._tc.iter(qn("w:sdt")):
        properties = control.find(qn("w:sdtPr"))
        checkbox = properties.find(qn("w14:checkbox")) if properties is not None else None
        if checkbox is None:
            continue
        checked = checkbox.find(qn("w14:checked"))
        if checked is None:
            checked = checkbox.makeelement(qn("w14:checked"))
            checkbox.insert(0, checked)
        checked.set(qn("w14:val"), "1")
        content = control.find(qn("w:sdtContent"))
        if content is not None:
            for text_node in content.iter(qn("w:t")):
                text_node.text = "☑"
        return True
    return False


def _append_check(cell: Any, boxed: bool = False) -> None:
    if boxed and _check_word_checkbox(cell):
        return
    if boxed:
        for paragraph in cell.paragraphs:
            if "□" in paragraph.text:
                _replace_box_in_paragraph(paragraph, paragraph.text.find("□"))
                return
    cell.paragraphs[-1].add_run("  √")


def _replace_box_in_paragraph(paragraph: Any, offset: int) -> bool:
    cursor = 0
    for run in paragraph.runs:
        next_cursor = cursor + len(run.text)
        if cursor <= offset < next_cursor:
            run_offset = offset - cursor
            run.text = run.text[:run_offset] + "☑" + run.text[run_offset + 1:]
            return True
        cursor = next_cursor
    return False


def _mark_option(cell: Any, label: str) -> None:
    for paragraph in getattr(cell, "paragraphs", [cell]):
        marker_offset = paragraph.text.find("□" + label)
        if marker_offset >= 0:
            _replace_box_in_paragraph(paragraph, marker_offset)
            return


def _gender_label(value: Any) -> str:
    return {"male": "男", "female": "女", "男": "男", "女": "女"}.get(str(value).lower(), str(value or "未填写"))


def _fill_inflammation_values(cell: Any, crp: Any, il6: Any) -> None:
    values = [value for value in (crp, il6) if value is not None]
    if not values:
        return
    for paragraph in cell.paragraphs:
        if "CRP" not in paragraph.text or "IL-6" not in paragraph.text:
            continue
        value_index = 0
        for run in paragraph.runs:
            if value_index >= len(values):
                break
            if run.font.underline and not run.text.strip():
                run.text = f"  {values[value_index]}  "
                value_index += 1
        return


def _set_cell(cell: Any, value: Any) -> None:
    cell.text = str(value)


def _output_reference(output_path: Path) -> str:
    try:
        return str(output_path.relative_to(SCALES_DIR.parent))
    except ValueError:
        return str(output_path)


def _set_inline_patient_info(paragraph: Any, payload: Any) -> None:
    paragraph.clear()
    paragraph.add_run(
        f"姓名：{_value(payload, 'patient_name') or '未填写'}    "
        f"性别：{_gender_label(_value(payload, 'gender'))}    "
        f"年龄：{_value(payload, 'age')}岁    "
        "病区：____    床号：____    住院号：____"
    )


def generate_nrs_document(payload: Any, result: Dict[str, Any]) -> str:
    document = Document(TEMPLATE_FILES["nrs"])
    info, score_table = document.tables
    _set_cell(info.cell(1, 1), _value(payload, "patient_name") or "未填写")
    _set_cell(info.cell(1, 4), _gender_label(_value(payload, "gender")))
    _set_cell(info.cell(1, 6), f"{_value(payload, 'age')}岁")
    _set_cell(info.cell(2, 1), _value(payload, "height"))
    _set_cell(info.cell(2, 4), _value(payload, "weight_records")["0"])
    _set_cell(info.cell(2, 6), result["bmi"])
    _append_check(score_table.cell(2 + int(result["nutrition_score"]), 3))
    _append_check(score_table.cell(6 + int(result["disease_score"]), 3))
    _append_check(score_table.cell(10 + int(result["age_score"]), 3))
    _set_cell(score_table.cell(12, 3), result["total_score"])
    _set_cell(score_table.cell(13, 3), datetime.now().strftime("%Y-%m-%d"))
    output_path = _new_output_path("NRS-2002", payload)
    document.save(output_path)
    return _output_reference(output_path)


def generate_mna_document(payload: Any, result: Dict[str, Any]) -> str:
    document = Document(TEMPLATE_FILES["mna"])
    _set_inline_patient_info(document.paragraphs[2], payload)
    table = document.tables[0]
    scores = result["score_breakdown"]
    for row, key in enumerate(("q1_appetite", "q2_weight_loss", "q3_mobility", "q4_stress", "q5_mental"), start=1):
        _append_check(table.cell(row, int(scores[key]) + 1), boxed=True)
    q6_row = 6 if result.get("q6_scoring_method") != "calf" else 7
    _append_check(table.cell(q6_row, int(scores["q6_bmi_or_calf"]) + 1), boxed=True)
    _set_cell(table.cell(8, 1), result["total_score"])
    _set_cell(table.cell(9, 1), datetime.now().strftime("%Y-%m-%d"))
    output_path = _new_output_path("MNA-SF", payload)
    document.save(output_path)
    return _output_reference(output_path)


def generate_glim_document(payload: Any, result: Dict[str, Any]) -> str:
    document = Document(TEMPLATE_FILES["glim"])
    _set_inline_patient_info(document.paragraphs[2], payload)
    phenotypic, etiological = document.tables
    records = _value(payload, "weight_records")
    age = int(_value(payload, "age"))
    bmi = float(result["bmi"])
    if "非自主体重丢失" in result["phenotypic_criteria_triggered"]:
        _mark_option(phenotypic.cell(1, 0), "非自主体重丢失")
        if result["weight_loss_6m_pct"] is not None and result["weight_loss_6m_pct"] > 5:
            _mark_option(phenotypic.cell(1, 1), "6个月内体重丢失＞5%")
        if result["weight_loss_over6m_pct"] is not None and result["weight_loss_over6m_pct"] > 10:
            _mark_option(phenotypic.cell(1, 1), "6个月以上体重丢失＞10%")
    if "低BMI" in result["phenotypic_criteria_triggered"]:
        _mark_option(phenotypic.cell(2, 0), "低BMI")
        _mark_option(phenotypic.cell(2, 1), "70岁以下BMI＜18.5 kg/m2" if age < 70 else "70岁以上BMI＜20 kg/m2")
    if "肌肉减少" in result["phenotypic_criteria_triggered"]:
        _mark_option(phenotypic.cell(3, 0), "肌肉减少")
        _mark_option(phenotypic.cell(3, 1), "人体成分分析提示肌肉减少，目前缺乏统一的切点值")

    if any((_value(payload, "intake_under_50_over_1w"), _value(payload, "any_intake_reduction_over_2w"), _value(payload, "gi_symptoms"), _value(payload, "nutrition_impact_conditions"))):
        _mark_option(etiological.cell(1, 0), "摄食减少或")
    if _value(payload, "intake_under_50_over_1w"):
        _mark_option(etiological.cell(1, 1), "摄入量≤50%的能量需求超过一周")
    if _value(payload, "any_intake_reduction_over_2w"):
        _mark_option(etiological.cell(1, 1), "任何摄入量减少超过2周")
    for item in _value(payload, "gi_symptoms", []):
        if item in GI_LABELS:
            _mark_option(etiological.cell(1, 1), GI_LABELS[item])
    for item in _value(payload, "nutrition_impact_conditions", []):
        if item in NUTRITION_IMPACT_LABELS:
            _mark_option(etiological.cell(1, 2), NUTRITION_IMPACT_LABELS[item])

    has_inflammation = bool(_value(payload, "acute_disease_ids", []) or _value(payload, "chronic_disease_ids", []) or (result.get("etiological_met") and ((_value(payload, "crp") is not None and _value(payload, "crp") >= 5) or (_value(payload, "il6") is not None and _value(payload, "il6") >= 7))))
    if has_inflammation:
        _mark_option(etiological.cell(2, 0), "炎症或疾病负担")
    for item in _value(payload, "acute_disease_ids", []):
        if item in ACUTE_DISEASE_LABELS:
            _mark_option(etiological.cell(2, 1), ACUTE_DISEASE_LABELS[item])
    for item in _value(payload, "chronic_disease_ids", []):
        if item in CHRONIC_DISEASE_LABELS:
            _mark_option(etiological.cell(2, 2), CHRONIC_DISEASE_LABELS[item])
    if "malignant_tumor" in _value(payload, "chronic_disease_ids", []):
        details = []
        if _value(payload, "cancer_site"):
            details.append(f"具体部位：{_value(payload, 'cancer_site')}")
        if _value(payload, "cancer_stage"):
            details.append(f"癌症分期：{_value(payload, 'cancer_stage')}")
        related = _value(payload, "cancer_malnutrition_related", None)
        if related is not None:
            _mark_option(etiological.cell(2, 2), "是" if related else "否")
        if details:
            etiological.cell(2, 2).add_paragraph("；".join(details))

    if _value(payload, "crp") is not None or _value(payload, "il6") is not None:
        marker_cell = etiological.cell(3, 1)
        if (_value(payload, "crp") is not None and _value(payload, "crp") >= 5) or (_value(payload, "il6") is not None and _value(payload, "il6") >= 7):
            _mark_option(marker_cell, "炎症状态指标")
        _fill_inflammation_values(marker_cell, _value(payload, "crp"), _value(payload, "il6"))

    if result["is_malnourished"] and result["severity"]:
        _mark_option(document.paragraphs[8], "中度营养不良" if "中度" in result["severity"] else "重度营养不良")
    output_path = _new_output_path("GLIM", payload)
    document.save(output_path)
    return _output_reference(output_path)


def _mark_sga_grade(table: Any, row: int, grade: int) -> None:
    _append_check(table.cell(row, grade + 1), boxed=True)


def _sga_grade_from_weight(records: Dict[str, Any]) -> Any:
    if not records.get("0") or not records.get("6"):
        return None
    current = float(records["0"])
    six_months = float(records["6"])
    loss = (six_months - current) / six_months * 100
    if loss > 10 and records.get("1") and current > float(records["1"]):
        return 0
    if loss < 5:
        return 0
    if loss <= 10:
        return 1
    return 2


def _sga_overall_result(grades: list[Any]) -> str:
    if any(grade is None for grade in grades):
        return "SGA评估信息不足"
    c_count = sum(grade == 2 for grade in grades)
    bc_count = sum(grade >= 1 for grade in grades)
    if c_count >= 5:
        return "SGA-C级"
    if bc_count >= 5:
        return "SGA-B级"
    return "SGA-A级"


def sga_overall_result(draft: Dict[str, Any], result: Dict[str, Any]) -> str:
    """Compatibility wrapper; the SGA service is the sole grading authority."""
    return sga_service.evaluate_sga(draft, result)["code"]


def _fill_sga_overall_result(document: Any, value: str) -> None:
    for paragraph in document.paragraphs:
        if "营养状况评估结果：" not in paragraph.text:
            continue
        for run in paragraph.runs:
            if run.font.underline and not run.text.strip():
                run.text = value
                return
        paragraph.add_run(value)
        return


def generate_sga_document(draft: Dict[str, Any], result: Dict[str, Any]) -> str:
    patient = draft.get("patient_info") or {}
    payload = {"patient_name": patient.get("name", ""), "gender": patient.get("gender", ""), "age": patient.get("age", "")}
    document = Document(TEMPLATE_FILES["sga"])
    header = document.paragraphs[1]
    header.clear()
    header.add_run(f"科室名称：____    姓名：{payload['patient_name'] or '未填写'}    病历号：____    床位：____    日期：{datetime.now().strftime('%Y-%m-%d')}")
    table = document.tables[0]
    records = draft.get("weight_records") or {}
    weight_grade = _sga_grade_from_weight(records)
    if weight_grade is not None:
        _mark_sga_grade(table, 1, weight_grade)

    intake_grade = {100: 0, 75: 1, 50: 1, 25: 2}.get(int(float((draft.get("intake_records") or {}).get("4") or -1)))
    if intake_grade is not None:
        _mark_sga_grade(table, 2, intake_grade)

    image_form = draft.get("image_screening_form") or {}
    gi_grade = {"none": 0, "mild_under_2w": 1, "severe_over_2w": 2}.get(image_form.get("giSymptoms"))
    if gi_grade is not None:
        _mark_sga_grade(table, 3, gi_grade)

    mna_form = draft.get("mnasf_form") or {}
    mobility_value = mna_form.get("mobility")
    if mobility_value is None:
        mobility_value = ((draft.get("mnasf_result") or {}).get("score_breakdown") or {}).get("q3_mobility")
    mobility_grade = {"2": 0, "1": 1, "0": 2}.get(str(mobility_value))
    if mobility_grade is not None:
        _mark_sga_grade(table, 4, mobility_grade)

    stress_grade = {"no_fever": 0, "temp_37_to_39_3d": 1, "temp_ge_39_over_3d": 2}.get(image_form.get("stressResponse"))
    if stress_grade is not None:
        _mark_sga_grade(table, 5, stress_grade)

    normalized_score = float(result["sga_normalized_score"])
    facial_grade = 0 if normalized_score < 3.5 else 1 if normalized_score < 6.3 else 2
    _mark_sga_grade(table, 6, facial_grade)
    _mark_sga_grade(table, 7, facial_grade)

    edema_grade = {"none": 0, "mild_moderate": 1, "severe": 2}.get(image_form.get("ankleEdema"))
    if edema_grade is not None:
        _mark_sga_grade(table, 8, edema_grade)

    _fill_sga_overall_result(document, sga_service.evaluate_sga(draft, result)["code"])

    output_path = _new_output_path("SGA", payload)
    document.save(output_path)
    return _output_reference(output_path)
